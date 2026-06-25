from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
IMG_DIR = ROOT / "review_work" / "images"
OUT = ROOT / "EA_Final_Review_Guide_CN.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    for r in p.runs:
        r.font.name = "Calibri"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def set_doc_fonts(doc):
    styles = doc.styles
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(10.5)
    styles["Heading 1"].font.size = Pt(18)
    styles["Heading 2"].font.size = Pt(14)
    styles["Heading 3"].font.size = Pt(12)


def add_title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Enterprise Architecture 期末复习资料")
    r.bold = True
    r.font.size = Pt(22)
    r.font.name = "Calibri"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Lecture 1-6 + Seminar 1-7 核心知识点讲解版")
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(80, 80, 80)

    doc.add_paragraph(
        "使用方法：先看每讲“主线”，再背“必考点”和“易混点”；带有【Discussion】标记的内容优先级更高。"
        "英文术语保留在正文里，是为了直接对应选择题原文。"
    )


def add_box(doc, title, lines, fill="EAF3FF"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(11)
    for line in lines:
        p = cell.add_paragraph(line, style=None)
        p.paragraph_format.space_after = Pt(2)
    doc.add_paragraph()


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_image(doc, glob_name, caption):
    matches = sorted(IMG_DIR.glob(glob_name))
    if not matches:
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(matches[0]), width=Inches(5.9))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    cap.runs[0].font.size = Pt(9)


def add_term_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True)
        set_cell_shading(hdr[i], "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val)
    doc.add_paragraph()


def add_lecture(doc, num, title, source, mainline, must_know, how_to_test, discussion, mistakes, images):
    doc.add_heading(f"Lecture {num}: {title}", level=1)
    doc.add_paragraph(f"来源：{source}")
    add_box(doc, "这一讲的主线", [mainline], fill="F3F7FB")
    doc.add_heading("必背知识点", level=2)
    add_bullets(doc, must_know)
    if images:
        for glob_name, caption in images:
            add_image(doc, glob_name, caption)
    doc.add_heading("选择题常见问法", level=2)
    add_bullets(doc, how_to_test)
    doc.add_heading("【Discussion】优先复习", level=2)
    add_bullets(doc, discussion)
    doc.add_heading("易混点", level=2)
    add_bullets(doc, mistakes)


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)
    set_doc_fonts(doc)
    add_title(doc)

    doc.add_heading("考前总览", level=1)
    add_box(
        doc,
        "最高优先级",
        [
            "1. Discussion slides：老师明确说会从 Lecture 的 Discussion 出题，看到 Discussion 就当成重点。",
            "2. 模型/框架的“用途”和“组成”：选择题常考某个模型是干什么的、包括哪些部分。",
            "3. 易混英文术语：alignment、governance、framework、model、view、viewpoint、stakeholder、concern。",
            "4. Seminar 应用：BSC、BMC、EFQM、conceptual model、Zachman first row、relational data model、ArchiMate layers。",
        ],
        fill="FFF3D6",
    )
    add_term_table(
        doc,
        ["考试词", "中文理解", "最容易考的点"],
        [
            ("Enterprise Architecture (EA)", "企业架构", "用原则、方法、模型把组织结构、流程、信息系统和技术基础设施连起来。"),
            ("Alignment", "对齐/一致性", "business 与 IT、strategy 与 people、process 与 customers 的对应关系。"),
            ("Framework", "框架", "用来分类和组织架构内容；Zachman 更像分类表，TOGAF 更像方法流程。"),
            ("Model", "模型", "对现实的一种抽象表达，不是现实本身。"),
            ("View", "视图", "给某类 stakeholder 看的模型表达。"),
            ("Viewpoint", "视点/视角模板", "规定某类 view 应该怎样表达、回答什么 concern。"),
            ("Stakeholder", "利益相关者", "对系统/企业有 concern 的人或组织。"),
            ("Concern", "关注点", "stakeholder 关心的问题，例如成本、风险、安全、效率。"),
        ],
    )

    add_lecture(
        doc,
        1,
        "EA introduction and basic concepts",
        "Lecture 1, Seminar 1",
        "第一讲是在建立 EA 的基本语言：什么是 architecture、enterprise、stakeholder，以及为什么企业需要 EA。",
        [
            "Architecture 的核心不是“画图”，而是描述一个系统在环境中的 fundamental concepts/properties、elements、relationships、principles。",
            "Enterprise 指拥有共同目标或共同 bottom line 的组织集合，不一定只是一家公司，也可以是部门、政府机构、合作网络。",
            "Enterprise Architecture 是 principles、methods、models 的 coherent whole，用来设计和实现组织结构、业务流程、信息系统和技术基础设施。",
            "EA 的价值：控制复杂性、减少重复系统、提高 business-IT alignment、支持沟通、支持治理和决策。",
            "Stakeholder 是对企业/系统有关注点的人或组织；concern 是他们真正关心的问题，例如成本、合规、性能、数据质量。",
            "EA 不等于 IT architecture。EA 范围更大，包括 business architecture、information/data architecture、application architecture、technology architecture。",
        ],
        [
            "问 definition：EA 是 coherent whole of principles, methods and models，而不是单纯的软件架构或数据库设计。",
            "问 stakeholder：答案通常是“has concerns/interests in the system”。",
            "问为什么需要 EA：选择“bridge communication gap / manage complexity / improve alignment”的选项。",
        ],
        [
            "What is stakeholder? 重点答法：stakeholder = 对系统有利益或关注点的人/组织；例子：manager 关心成本，customer 关心服务质量，IT staff 关心可维护性。",
            "Why do we need Enterprise Architecture? 重点答法：因为企业越来越复杂，业务和 IT 容易脱节，EA 用模型和原则让不同角色对同一套结构达成理解。",
        ],
        [
            "Enterprise 不等于 information system；企业包括组织、流程、人员、数据、应用和技术。",
            "Architecture model 不是最终目的，模型是为了沟通、分析和决策。",
        ],
        [],
    )

    add_lecture(
        doc,
        2,
        "EA process, drivers, alignment and governance instruments",
        "Lecture 2, Seminar 2",
        "第二讲把 EA 放进企业管理场景：企业为什么要做 EA，EA 如何帮助战略、流程、客户、IT 互相对齐。",
        [
            "Architecture process 是从 initial idea 经过 design、implementation，到 operational system，再到 changing/replacing system 的闭环过程；所有阶段都需要 stakeholder 之间清晰沟通。",
            "Internal drivers 来自组织内部，例如战略改变、流程复杂、系统重复、成本压力、业务与 IT 不一致。",
            "External drivers 来自外部环境，例如法律法规、竞争、市场变化、技术变化、客户需求变化。",
            "Strategic Alignment Model 关注 business strategy、IT strategy、organizational infrastructure/processes、IT infrastructure/processes 之间的一致性。",
            "Nadler 的 congruence/alignment 组件：work、people、formal organisation、informal organisation。",
            "Vertical alignment = top strategy 与 bottom people/action 的关系；Horizontal alignment = internal processes 与 external customers 的关系。",
            "Operating model 根据 business process integration 和 business process standardization 两个维度理解企业运作。",
            "Balanced Scorecard (BSC) 把 vision/strategy 翻译成 objectives、measures、targets、initiatives；常见四视角是 financial、customer、internal business process、learning and growth。",
            "Business Model Canvas (BMC) 是高层业务模式模板，通常包括 customer segments、value propositions、channels、customer relationships、revenue streams、key resources、key activities、key partners、cost structure。",
        ],
        [
            "给出 horizontal/vertical alignment，让你判断分别对应什么。",
            "问 drivers：internal 是企业内部问题，external 是企业外部压力。",
            "问 BSC 的目的：不是画商业模式，而是把战略落成可衡量目标。",
            "问 BMC 的组成：看到 value propositions、customer segments、revenue streams 等就是 BMC。",
        ],
        [
            "Describe internal and external drivers for your business/project：答题要分两栏，不要混在一起。",
            "Good enterprise architecture characteristics：一般围绕 clear、consistent、understandable、business-aligned、communicable、usable for decision-making。",
        ],
        [
            "BSC 和 BMC 不一样：BSC 管战略执行和绩效衡量；BMC 管业务模式怎么创造和获取价值。",
            "Vertical alignment 不是 internal process 到 customer；那是 horizontal alignment。",
        ],
        [
            ("L2_p07_strategic_alignment*.png", "图 1：Strategic Alignment Model，理解 business strategy 与 IT strategy 的对齐。"),
            ("L2_p11_operating_model*.png", "图 2：Operating model，用 integration 与 standardization 理解企业运作方式。"),
            ("L2_p25_business_model_canvas*.png", "图 3：Business Model Canvas 九宫格，Seminar 2/7 常用。"),
        ],
    )

    add_lecture(
        doc,
        3,
        "Governance instruments, COBIT, Zachman, TOGAF and MDA",
        "Lecture 3, Seminar 3, Seminar 5",
        "第三讲内容最多，考试也最容易出选择题：每个治理工具或架构框架要记住“用途、组成、区别”。",
        [
            "EFQM 是质量/卓越管理模型，范围比 ISO 9001 更宽，不只关注质量管理，还提供整体 performance improvement 框架。",
            "COBIT = Control Objectives for Information and Related Technology，是 IT governance/control framework，桥接 business risks、control needs、technical issues。",
            "COBIT 五个 process domains：Evaluate/Direct/Monitor (EDM)、Align/Plan/Organise (APO)、Build/Acquire/Implement (BAI)、Deliver/Service/Support (DSS)、Monitor/Evaluate/Assess (MEA)。",
            "COBIT maturity levels：Ad Hoc、Repeatable、Defined、Managed、Optimised。题目常考顺序和每级特征。",
            "ITIL 与 COBIT 互补：COBIT tells what to do；ITIL explains how to do it。",
            "CMMI 用来判断软件工程/过程能力成熟度；常见等级包括 Initial、Managed、Defined、Quantitatively Managed、Optimizing。",
            "IEEE 1471/ISO 42010 标准强调 stakeholder、concern、view、viewpoint、architecture description 的关系。",
            "Zachman Framework 是分类框架，用 6 个问题 What/How/Where/Who/When/Why 与不同视角组织架构产物；优点是易理解、覆盖企业整体，缺点是关系和方法不够具体。",
            "TOGAF 是框架加方法，核心是 ADM，一个迭代循环，用来开发和管理企业架构。",
            "TOGAF 主要组件包括 Architecture Capability Framework、ADM、Architecture Content Framework、Enterprise Continuum。",
            "MDA = Model-Driven Architecture，通过 CIM/PIM/PSM 等抽象层和 mapping，从模型走向实现。",
        ],
        [
            "问 COBIT 五域：看到 Build, acquire and implement 就选 COBIT domain。",
            "问 COBIT vs ITIL：COBIT 偏 governance/control objective，ITIL 偏 service management best practice。",
            "问 Zachman vs TOGAF：Zachman 更像分类表，TOGAF 更像方法论和过程。",
            "问 ADM：TOGAF 的核心方法，是 iterative/cyclic。",
            "问 MDA：重点是模型驱动、抽象层、从 platform-independent 到 platform-specific。",
        ],
        [
            "EFQM vs ISO 9001：ISO 9001 更偏质量管理体系标准；EFQM 更宽，关注组织卓越和绩效改进。",
            "ITIL vs COBIT：COBIT 给控制目标和治理方向；ITIL 给服务管理实践。",
        ],
        [
            "不要把 COBIT 的 BAI 误认为 Business Model Canvas；BAI 是 Build/Acquire/Implement。",
            "Zachman 不是实施步骤；TOGAF ADM 才是具体开发流程。",
            "Optimised/Optimizing 表示持续改进，不是刚开始建立流程。",
        ],
        [
            ("L3_p12_cobit_domains*.png", "图 4：COBIT process domains，五个域要背英文。"),
            ("L3_p29_zachman_framework*.png", "图 5：Zachman Framework，用问题维度组织架构产物。"),
            ("L3_p39_togaf_components*.png", "图 6：TOGAF main components。"),
            ("L3_p41_togaf_adm*.png", "图 7：TOGAF ADM，核心是迭代循环。"),
            ("L3_p45_mda_levels*.png", "图 8：MDA 抽象层与映射。"),
        ],
    )

    add_lecture(
        doc,
        4,
        "Description languages: IDEF, BPMN and UML",
        "Lecture 4",
        "第四讲讲建模语言。考试不会要求你真的画复杂模型，主要考每种语言适合表达什么、缺点是什么。",
        [
            "IDEF 是一族 enterprise modelling and analysis languages。核心包括 IDEF0、IDEF3、IDEF1X。",
            "IDEF0 用于 functional modelling，关注一个 function/activity 如何被 input、control、mechanism 影响并产生 output。",
            "IDEF0 五件套：Activity、Input、Output、Control、Mechanism。口诀：AICOM。",
            "IDEF3 用于 process modelling，捕捉业务流程 workflow。",
            "IDEF1X 用于 data modelling，建立 logical/physical data models。",
            "IDEF 的问题：不同模型间 communication mechanisms 弱，模型孤立，不利于把架构系统整体可视化。",
            "BPMN 用于 business process modelling，提供统一业务流程符号；范围主要限制在流程，不覆盖完整 EA 的应用、数据和技术全景。",
            "UML 原本面向 software/system design，图类型多，业务人员不一定容易理解；后来也被扩展到架构建模。",
            "Suitability for EA 的核心结论：IDEF/BPMN/UML 都有价值，但都不是完整 EA 语言；跨域关系和整合通常不足。",
        ],
        [
            "问 IDEF0 五元素：Activity/Input/Output/Control/Mechanism。",
            "问 BPMN 适合什么：business process/workflow。",
            "问 IDEF1X 适合什么：data modelling。",
            "问这些语言为什么不完全适合 EA：因为跨视图、跨域关系表达不足。",
        ],
        [
            "TOGAF vs Zachman：一个偏方法过程，一个偏分类结构。",
            "BPMN vs UML vs IDEF：BPMN 画流程，UML 偏软件系统设计，IDEF 分功能/过程/数据几类建模。",
        ],
        [
            "IDEF0 的 Control 不是 Output；Control 是约束/规则/条件，Output 是活动产生的结果。",
            "BPMN 不等于完整 EA language，它主要解决 process modelling。",
        ],
        [
            ("L4_p10_idef0*.png", "图 9：IDEF0，Activity 周围的 Input/Output/Control/Mechanism。"),
            ("L4_p16_bpmn*.png", "图 10：BPMN 示例，重点看它表达流程。"),
        ],
    )

    add_lecture(
        doc,
        5,
        "Architecture descriptions, integration and model semantics",
        "Lecture 5",
        "第五讲解决一个现实问题：企业架构里有很多异构模型，怎样描述、整合、沟通和维护它们。",
        [
            "Enterprise architecture descriptions 经常包含许多 heterogeneous models，例如 UML、BPMN、Zachman cells 等。",
            "Integration of architectural domains 是关键难题。业务流程模型、应用模型、技术模型如果彼此孤立，就会产生 alignment problem。",
            "Business-IT alignment problem 的典型例子：业务流程变了，但应用系统或技术基础设施没有同步变化。",
            "Architecture description 是对架构的表达，可以包含 models、views、principles、decisions 等。",
            "Architecture language 要支持 five ways：way of thinking、way of modelling、way of communicating、way of working、way of supporting。",
            "Symbolic model 是符号/语法结构，例如用图形或形式化符号描述 Employee、Director、Responsible_for。",
            "Semantic model 是对 symbolic model 的解释，说明这些符号在现实或某个 universe 中对应什么。",
            "Subjective model/visualization 体现 stakeholder 对模型的理解或视图表达；同一底层模型可以有不同可视化。",
        ],
        [
            "问 alignment problem：模型/域之间缺少连接，业务和 IT 难以同步。",
            "问 architecture language 的作用：支持建模、沟通、分析、决策和工具化管理。",
            "问 symbolic vs semantic：symbolic 是符号结构，semantic 是意义解释。",
        ],
        [
            "How different enterprise architecture domains can be integrated? 重点围绕 common language、defined relationships、shared concepts、traceability。",
            "Symbolic/semantic/subjective models integration：不要只背术语，要理解为“符号、含义、人的视图”三个层次。",
        ],
        [
            "Model integration 不是把图贴在一起；要能表达它们之间的关系和影响。",
            "View 不是 viewpoint。View 是具体结果，viewpoint 是生成这种视图的规则/模板。",
        ],
        [
            ("L5_p13_observing_universe*.png", "图 11：观察 universe、stakeholder 与 architecture description 的关系。"),
        ],
    )

    add_lecture(
        doc,
        6,
        "Communication of EA and ArchiMate language",
        "Lecture 6, Seminar 6",
        "第六讲是建模语言重点：为什么需要统一语言，以及 ArchiMate 如何把业务、应用、技术和动机/实施层连起来。",
        [
            "EA communication 的目的不是展示复杂图，而是让 stakeholder 获得正确知识并达成理解、认可或承诺。",
            "常见沟通形式包括 workshop、validation interview、committing review、presentation、mailing。",
            "ArchiMate 是 enterprise architecture modelling language，用统一的概念和关系描述不同层的模型。",
            "ArchiMate premise：不同层的模型有相似的一般结构，因此可以用统一的概念类型和关系跨层建模。",
            "核心三层：Business layer、Application layer、Technology layer。",
            "Business layer 表示组织对外部客户提供的业务服务、业务流程、业务角色/actor 等。",
            "Application layer 表示应用组件、应用服务、应用接口、数据对象等，连接业务需求与技术实现。",
            "Technology layer 表示 node、device、system software、technology service、artifact 等技术基础设施。",
            "Motivation concepts 包括 stakeholder、driver、assessment、goal、outcome、principle、requirement、constraint、meaning、value。",
            "Strategy concepts 包括 resource、capability、course of action、value stream。",
            "Business-application alignment：应用层可以服务/支撑业务层，常通过服务关系和 realization/serving 等关系理解。",
            "Application-technology alignment：技术层承载或支撑应用层，例如 node/device 部署 artifact 或提供 technology service。",
            "Implementation/migration concepts 包括 work package、deliverable、implementation event、plateau、gap，用于表达变更项目和迁移路径。",
            "BMC 可以映射到 ArchiMate：客户、价值主张、关键活动、资源、渠道等可以分别映射到业务层、动机层或策略层概念。",
        ],
        [
            "问 ArchiMate 三层：Business/Application/Technology。",
            "问 common well-defined vocabulary 支持什么：architecture design support / communication / integration。",
            "问某个概念属于哪层：business actor 是 business layer；application component 是 application layer；node/device 是 technology layer。",
            "问 implementation concepts：work package、deliverable、plateau、gap。",
            "问 Motivation：goal、requirement、stakeholder、driver、assessment 等。",
        ],
        [
            "EA communication theories：关注 stakeholder、knowledge goal、conversation strategy、agreement/commitment。",
            "Business Model Canvas mapped to ArchiMate：重点不是死记论文细节，而是理解 BMC 的商业元素可以转成 ArchiMate 的业务/动机/策略模型。",
        ],
        [
            "Business service 不等于 application service：前者面向业务客户，后者由应用提供给业务或其他应用。",
            "Node/device 是技术层，不是应用层。",
            "Requirement 和 goal 都在 motivation 相关概念中，但 requirement 更具体，表示必须满足的需要。",
        ],
        [
            ("L6_p18_archimate_layers*.png", "图 12：ArchiMate 层次思想，不同层结构相似。"),
            ("L6_p20_archimate_structure*.png", "图 13：ArchiMate language structure。"),
            ("L6_p29_business_layer*.png", "图 14：Business layer model 示例。"),
            ("L6_p40_business_application_alignment*.png", "图 15：Business-Application alignment。"),
            ("L6_p42_technology_layer*.png", "图 16：Technology layer model 示例。"),
            ("L6_p50_implementation*.png", "图 17：Implementation/migration concepts。"),
        ],
    )

    doc.add_heading("Seminar 应用复习表", level=1)
    add_term_table(
        doc,
        ["Seminar", "你需要会做什么", "对应 Lecture/考试点"],
        [
            ("Seminar 1", "确定 business、customer、internal process、stakeholder 等基础信息。", "Lecture 1：stakeholder、enterprise、EA 基础。"),
            ("Seminar 2", "用 BSC 做战略目标/指标表，用 BMC 描述商业模式。", "Lecture 2：BSC vs BMC。"),
            ("Seminar 3", "基于 EFQM Excellence Model 分析案例。", "Lecture 3：EFQM、ISO 9001 区别。"),
            ("Seminar 4", "创建 conceptual model 并解释实体/关系。", "Lecture 3/5：model、architecture description。"),
            ("Seminar 5", "填写 Zachman first row，并把 conceptual model 转 relational data model。", "Lecture 3/4：Zachman、data modelling。"),
            ("Seminar 6", "用 ArchiMate 创建 business layer，application/technology layer 可加分。", "Lecture 6：ArchiMate layers。"),
            ("Seminar 7", "整合前六次 Seminar 报告。", "综合题：BSC、BMC、conceptual model、Zachman、relational model、ArchiMate。"),
        ],
    )

    doc.add_heading("Discussion 高频考法清单", level=1)
    add_bullets(
        doc,
        [
            "L1：为什么需要 EA？答复杂性、沟通鸿沟、业务 IT 对齐、治理和决策。",
            "L1：stakeholder/enterprise 的含义和例子。",
            "L2：internal vs external drivers；good EA 的 characteristics。",
            "L3：ISO 9001 vs EFQM；ITIL vs COBIT。",
            "L4：TOGAF vs Zachman；BPMN vs UML vs IDEF。",
            "L5：different EA domains 如何 integrated；symbolic/semantic/subjective models 如何理解。",
            "L6：EA communication theories；BMC 如何 mapped to ArchiMate。",
        ],
    )

    doc.add_heading("最后冲刺：最容易混的 20 个判断", level=1)
    add_numbered(
        doc,
        [
            "Vertical alignment = top strategy 到 bottom people/action。",
            "Horizontal alignment = internal processes 到 external customers。",
            "BSC = strategy execution and performance measurement。",
            "BMC = business model/value creation template。",
            "COBIT = IT governance/control framework。",
            "ITIL = IT service management best practices。",
            "COBIT tells what to do；ITIL explains how to do it。",
            "COBIT 五域包含 BAI = Build, Acquire and Implement。",
            "Zachman = classification schema/framework，不是实施方法。",
            "TOGAF ADM = iterative method/process，是 TOGAF 核心。",
            "MDA = model-driven, CIM/PIM/PSM abstraction and mapping。",
            "IDEF0 = functional modelling, AICOM 五元素。",
            "IDEF3 = process workflow。",
            "IDEF1X = data modelling。",
            "BPMN = business process modelling，不覆盖完整 EA。",
            "UML = software/system design language，业务人员未必易懂。",
            "Symbolic model = syntax/sign structure；semantic model = interpretation/meaning。",
            "ArchiMate 三层 = Business/Application/Technology。",
            "Business actor 属于 business layer；application component 属于 application layer；node/device 属于 technology layer。",
            "Work package、deliverable、plateau、gap 属于 implementation/migration concepts。",
        ],
    )

    doc.add_heading("建议背诵顺序", level=1)
    add_bullets(
        doc,
        [
            "第一轮：只背“最后冲刺 20 个判断”和 Discussion 高频清单。",
            "第二轮：按 Lecture 2、3、6 优先复习，因为选择题模型最多。",
            "第三轮：把 IDEF0、COBIT 五域、ArchiMate 三层和 Zachman/TOGAF 区别默写一遍。",
            "做题时遇到超长选项不要被干扰，先看关键词属于哪个模型或层。",
        ],
    )

    doc.add_paragraph("图片说明：文中插图均由本地 Lecture PDF 课件页截图生成，用于考前复习定位。")
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
